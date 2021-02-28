from operator import itemgetter
from itertools import cycle
from docx import Document
import warnings
import fitz
import json
import re


document_save = Document()
chapter = []
keywords = input("Enter a keyword to detect the pages for new chapter: ")

def fonts(doc, granularity=False):
    """Extracts fonts and their usage in PDF documents.
    :param doc: PDF document to iterate through
    :type doc: <class 'fitz.fitz.Document'>
    :param granularity: also use 'font', 'flags' and 'color' to discriminate text
    :type granularity: bool
    :rtype: [(font_size, count), (font_size, count}], dict
    :return: most used fonts sorted by count, font style information
    """
    styles = {}
    font_counts = {}
    page_number = 0

    for page in doc:
        page_number = page_number + 1
        blocks = page.getText("dict")["blocks"]
        for b in blocks:  # iterate through the text blocks
            if b['type'] == 0:  # block contains text
                for l in b["lines"]:  # iterate through the text lines
                    for s in l["spans"]:  # iterate through the text spans
                        if keywords in s['text']:
                            try:
                                chapter.append(page_number)
                                # print(s)
                            except Exception as e:
                                print(e)
                        if granularity:
                            identifier = "{0}_{1}_{2}_{3}".format(s['size'], s['flags'], s['font'], s['color'])
                            styles[identifier] = {'size': s['size'], 'flags': s['flags'], 'font': s['font'],
                                               'color': s['color']}
                        else:
                            identifier = "{0}".format(s['size'])
                            styles[identifier] = {'size': s['size'], 'font': s['font']}

                        font_counts[identifier] = font_counts.get(identifier, 0) + 1  # count the fonts usage

    font_counts = sorted(font_counts.items(), key=itemgetter(1), reverse=True)

    if len(font_counts) < 1:
        raise ValueError("Zero discriminating fonts found!")

    return font_counts, styles


def font_tags(font_counts, styles):
    """Returns dictionary with font sizes as keys and tags as value.
    :param font_counts: (font_size, count) for all fonts occuring in document
    :type font_counts: list
    :param styles: all styles found in the document
    :type styles: dict
    :rtype: dict
    :return: all element tags based on font-sizes
    """
    p_style = styles[font_counts[0][0]]  # get style for most used font by count (paragraph)
    p_size = p_style['size']  # get the paragraph's size

    # sorting the font sizes high to low, so that we can append the right integer to each tag
    font_sizes = []
    for (font_size, count) in font_counts:
        font_sizes.append(float(font_size))
    font_sizes.sort(reverse=True)

    # aggregating the tags for each font size
    idx = 0
    size_tag = {}
    for size in font_sizes:
        idx += 1
        if size == p_size:
            idx = 0
            size_tag[size] = '<p>'
        if size > p_size:
            size_tag[size] = '<pdftodocH{0}>'.format(idx)
        elif size < p_size:
            size_tag[size] = '<s{0}>'.format(idx)

    return size_tag


def headers_para(doc, size_tag,starting, ending):
    """Scrapes headers & paragraphs from PDF and return texts with element tags.
    :param doc: PDF document to iterate through
    :type doc: <class 'fitz.fitz.Document'>
    :param size_tag: textual element tags for each size
    :type size_tag: dict
    :rtype: list
    :return: texts with pre-pended element tags
    """
    header_para = []  # list with headers and paragraphs
    first = True  # boolean operator for first header
    previous_s = {}  # previous span

    for page in doc.pages(starting,ending,1): #64
        blocks = page.getText("dict")["blocks"]
        for b in blocks:  # iterate through the text blocks
            if b['type'] == 0:  # this block contains text

                # REMEMBER: multiple fonts and sizes are possible IN one block

                block_string = ""  # text found in block
                for l in b["lines"]:  # iterate through the text lines
                    for s in l["spans"]:  # iterate through the text spans
                        if s['text'].strip():  # removing whitespaces:
                            if first:
                                previous_s = s
                                first = False
                                block_string = size_tag[s['size']] + s['text']
                            else:
                                if s['size'] == previous_s['size']:

                                    if block_string and all((c == "|") for c in block_string):
                                        # block_string only contains pipes
                                        block_string = size_tag[s['size']] + s['text']
                                    if block_string == "":
                                        # new block has started, so append size tag
                                        block_string = size_tag[s['size']] + s['text']
                                    else:  # in the same block, so concatenate strings
                                        block_string += " " + s['text']

                                else:
                                    header_para.append(block_string)
                                    block_string = size_tag[s['size']] + s['text']

                                previous_s = s

                    # new block started, indicating with a pipe
                    block_string += ""
                header_para.append(block_string)

    return header_para


def main():

    with open("words.txt", "r") as words_file:
        finding_words = words_file.read().splitlines()
    finding_words = [" "+each_string.lower()+" " for each_string in finding_words]

    with open("ignoreword.txt", "r") as i_words_file:
        ignore_words = i_words_file.read().splitlines()
    ignore_words = [" " + each_string + " " for each_string in ignore_words]

    header_font_list = []
    document = 'EntireFIle.pdf'
    doc = fitz.open(document)
    # final = ""
    # i = 0
    # elements = []
    enter = True
    Split_File = input('Do you want to split output file into chapters? Please type 1 for yes otherwise type 0 and press enter : ')
    split = int(Split_File)
    print("Please type the tags of header fonts you want to add and type quit")
    print("Example : <pdftodocH17> \nYou can get the list of header tags from detect_header python script.")
    while enter:
        add_header_font = input('Type Header tag or type quit: ')
        if "<pdftodocH" in add_header_font:
            header_font_list.append(add_header_font)
        if "quit" in add_header_font:
            enter = False
    unique_list = list(set(header_font_list))
    #add_subscript_font = input('Please type the number of subscript font of your pdf file (Only one font allowed): ')

    font_counts, styles = fonts(doc, granularity=False)
    size_tag = font_tags(font_counts, styles)
    page_count = doc.pageCount
    print(chapter)

    if split ==0:
        final = ""
        i = 0
        elements = []
        document_save = Document()
        elements = headers_para(doc, size_tag, 20, 50)
        print("Total lines are " +str(len(elements)))

        for each_element in elements:
            if "<pdftodocH" in each_element:
                if final != "":
                    try:
                        string_list = final.split(". ")
                        #string_list = re.split(r'.(?=. [A-Z])', final)
                        for sk in string_list:
                            sk = sk+".\n"
                            for word in finding_words:
                                if word in sk:
                                    first ,last= sk.split(word)
                                    p = document_save.add_paragraph(first)
                                    runner = p.add_run(word)
                                    runner.bold = True
                                    p.add_run(last)
                                    break

                        document_save.save("demo.docx")
                    except Exception as e:
                        print(e)
                        print(each_element)
                for header_fonts in unique_list:
                    if  header_fonts in each_element:
                        try:
                            each_element = each_element.replace(header_fonts, "")
                            document_save.add_heading(each_element.replace(keywords,""))
                            document_save.save("demo.docx")
                        except Exception as e:
                            print(e)
                            print(each_element)
                final = ""

            if  "<p>" in each_element:
                each_element = each_element.replace("\n", " ")
                each_element = each_element.replace("<p>", " ")
                final = final+each_element

            elif  "<s228>" in each_element:
                each_element = each_element.replace("\n"," ")
                #final = final + each_element

            print(i," Lines Done")
            i = i+1

    elif split == 1:
        final = ""
        i = 0
        new_cycle = cycle(chapter)
        ending = next(new_cycle)
        running = True

        while running:
            document_save = Document()
            starting, ending = chapter[0], chapter[1]
            elements_ = []
            elements_ = headers_para(doc, size_tag, starting-1, ending-1)
            print("Total lines are " + str(len(elements_)))

            for each_element in elements_:
                if "<pdftodocH" in each_element:
                    if final != "":
                        try:
                            string_list = final.split(". ")
                            # string_list = re.split(r'.(?=. [A-Z])', final)
                            for sk in string_list:
                                sk = sk + ".\n"
                                for word in finding_words:
                                    if word in sk:
                                        first, last = sk.split(word)
                                        p = document_save.add_paragraph(first)
                                        runner = p.add_run(word)
                                        runner.bold = True
                                        p.add_run(last)
                                        break

                            document_save.save("Folder/"+"chapter_"+str(starting)+".docx")
                        except Exception as e:
                            print(e)
                            print(each_element)
                    for header_fonts in unique_list:
                        if header_fonts in each_element:
                            try:
                                each_element = each_element.replace(header_fonts, " ")
                                document_save.add_heading(each_element)
                                document_save.save("Folder/"+"chapter_"+str(starting)+".docx")
                            except Exception as e:
                                print(e)
                                print(each_element)
                    final = ""

                if "<p>" in each_element:
                    each_element = each_element.replace("\n", " ")
                    each_element = each_element.replace("<p>", " ")
                    final = final + each_element

                elif "<s228>" in each_element:
                    each_element = each_element.replace("\n", " ")
                    each_element = each_element.replace("<s228>", " ")
                    #final = final + each_element

                print(i, " Lines Done")
                i = i + 1
            chapter.pop(0)
            print()
            print(chapter)
            if ending == chapter[len(chapter) - 1]:
                print("Running False")
                running = False

if __name__ == '__main__':
    main()